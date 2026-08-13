"""Document parsing (§10).

Formats: HTML, PDF, XLSX/XLS, DOCX, TXT, JSON, CSV.

The rule that matters here is the second one in §10: **a financial table is not
flattened into prose.** A balance sheet turned into a stream of numbers teaches
the model that 3 698 502 257 and "итого активы" are unrelated tokens. Tables
are extracted as structured objects and rendered back as aligned Markdown, with
the header row preserved, so row/column association survives tokenisation.

Optional dependencies (pypdf, openpyxl, python-docx) are imported lazily. A
missing one downgrades that format to "unsupported" with a clear message; it
never crashes a dataset build that does not use it.
"""

from __future__ import annotations

import csv
import io
import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

SUPPORTED = ("html", "htm", "pdf", "xlsx", "xls", "docx", "txt", "json", "csv")


class UnsupportedFormat(RuntimeError):
    pass


@dataclass(slots=True)
class Table:
    header: list[str]
    rows: list[list[str]]
    caption: str | None = None
    page: int | None = None

    @property
    def shape(self) -> tuple[int, int]:
        return len(self.rows), len(self.header)

    def to_markdown(self) -> str:
        """Aligned Markdown. Column association is what we are preserving."""
        if not self.header and not self.rows:
            return ""
        header = self.header or [f"col{i+1}" for i in range(len(self.rows[0]))]
        widths = [len(h) for h in header]
        for row in self.rows:
            for index, cell in enumerate(row[: len(widths)]):
                widths[index] = max(widths[index], len(cell))
        def line(cells: list[str]) -> str:
            padded = [
                (cells[i] if i < len(cells) else "").ljust(widths[i])
                for i in range(len(widths))
            ]
            return "| " + " | ".join(padded) + " |"
        out = []
        if self.caption:
            out.append(self.caption.strip())
        out.append(line(header))
        out.append("|" + "|".join("-" * (w + 2) for w in widths) + "|")
        out.extend(line(row) for row in self.rows)
        return "\n".join(out)

    def as_dict(self) -> dict[str, Any]:
        return {
            "header": self.header,
            "rows": self.rows,
            "caption": self.caption,
            "page": self.page,
            "shape": list(self.shape),
        }


@dataclass(slots=True)
class ParsedDocument:
    text: str
    tables: list[Table] = field(default_factory=list)
    sections: list[dict[str, Any]] = field(default_factory=list)
    pages: int | None = None
    format: str = "txt"
    metadata: dict[str, Any] = field(default_factory=dict)

    def combined_text(self) -> str:
        """Prose plus rendered tables, in document order where known."""
        parts = [self.text.strip()] if self.text.strip() else []
        parts.extend(t.to_markdown() for t in self.tables if t.rows)
        return "\n\n".join(p for p in parts if p)


# --------------------------------------------------------------------------
# HTML
# --------------------------------------------------------------------------

_TABLE_RE = re.compile(r"<table\b.*?</table>", re.I | re.S)
_ROW_RE = re.compile(r"<tr\b.*?</tr>", re.I | re.S)
_CELL_RE = re.compile(r"<t([dh])\b[^>]*>(.*?)</t\1>", re.I | re.S)
_HEADING_RE = re.compile(r"<h([1-6])\b[^>]*>(.*?)</h\1>", re.I | re.S)
_CAPTION_RE = re.compile(r"<caption\b[^>]*>(.*?)</caption>", re.I | re.S)


def _cell_text(html: str) -> str:
    from ai.datasets.cleaning import normalize_encoding, strip_html

    return normalize_encoding(strip_html(html)).replace("\n", " ").strip()


def parse_html(html: str) -> ParsedDocument:
    from ai.datasets.cleaning import normalize_encoding, remove_navigation, strip_html

    tables: list[Table] = []
    for block in _TABLE_RE.findall(html):
        rows: list[list[str]] = []
        header: list[str] = []
        caption_match = _CAPTION_RE.search(block)
        for row_html in _ROW_RE.findall(block):
            cells = _CELL_RE.findall(row_html)
            values = [_cell_text(body) for _, body in cells]
            if not any(values):
                continue
            is_header = all(kind.lower() == "h" for kind, _ in cells) and cells
            if is_header and not header:
                header = values
            else:
                rows.append(values)
        if rows or header:
            tables.append(
                Table(
                    header=header,
                    rows=rows,
                    caption=_cell_text(caption_match.group(1)) if caption_match else None,
                )
            )

    sections = [
        {"level": int(level), "title": _cell_text(title)}
        for level, title in _HEADING_RE.findall(html)
        if _cell_text(title)
    ]
    prose = _TABLE_RE.sub("\n", html)
    text = remove_navigation(normalize_encoding(strip_html(prose)))
    return ParsedDocument(text=text, tables=tables, sections=sections, format="html")


# --------------------------------------------------------------------------
# PDF
# --------------------------------------------------------------------------

def parse_pdf(data: bytes) -> ParsedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise UnsupportedFormat(
            "PDF parsing needs pypdf: pip install -r ai/requirements.txt"
        ) from exc

    reader = PdfReader(io.BytesIO(data))
    chunks: list[str] = []
    tables: list[Table] = []
    for number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        chunks.append(f"[стр. {number}]\n{text.strip()}")
        tables.extend(_tables_from_layout(text, page=number))
    return ParsedDocument(
        text="\n\n".join(chunks),
        tables=tables,
        pages=len(reader.pages),
        format="pdf",
        metadata={"producer": (reader.metadata or {}).get("/Producer")},
    )


_NUMERIC = re.compile(r"^-?[\d\s.,()%-]+$")


def _tables_from_layout(text: str, *, page: int | None = None) -> list[Table]:
    """Recover column structure from a text-layer PDF.

    A block of consecutive lines that all split into the same number of
    whitespace-separated fields, at least two of which are numeric, is a table.
    This is a heuristic and it is honest about it: anything it is unsure about
    stays in the prose stream rather than becoming a fake table.
    """
    tables: list[Table] = []
    block: list[list[str]] = []

    def flush() -> None:
        if len(block) >= 3:
            header, *rows = block
            tables.append(Table(header=header, rows=rows, page=page))
        block.clear()

    previous_width = 0
    for line in text.splitlines():
        fields = [f for f in re.split(r"\s{2,}|\t", line.strip()) if f]
        numeric = sum(1 for f in fields if _NUMERIC.match(f))
        if len(fields) >= 3 and numeric >= 2:
            if previous_width and len(fields) != previous_width:
                flush()
            block.append(fields)
            previous_width = len(fields)
        else:
            flush()
            previous_width = 0
    flush()
    return tables


# --------------------------------------------------------------------------
# XLSX / XLS
# --------------------------------------------------------------------------

def parse_xlsx(path: str | Path) -> ParsedDocument:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover
        raise UnsupportedFormat(
            "XLSX parsing needs openpyxl: pip install -r ai/requirements.txt"
        ) from exc

    workbook = load_workbook(filename=str(path), read_only=True, data_only=True)
    tables: list[Table] = []
    for sheet in workbook.worksheets:
        rows = [
            ["" if cell is None else str(cell).strip() for cell in row]
            for row in sheet.iter_rows(values_only=True)
        ]
        rows = [r for r in rows if any(r)]
        if not rows:
            continue
        header, *body = rows
        tables.append(Table(header=header, rows=body, caption=sheet.title))
    workbook.close()
    return ParsedDocument(text="", tables=tables, format="xlsx")


# --------------------------------------------------------------------------
# DOCX
# --------------------------------------------------------------------------

def parse_docx(path: str | Path) -> ParsedDocument:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover
        raise UnsupportedFormat(
            "DOCX parsing needs python-docx: pip install -r ai/requirements.txt"
        ) from exc

    document = docx.Document(str(path))
    paragraphs: list[str] = []
    sections: list[dict[str, Any]] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = (paragraph.style.name or "").lower()
        if style.startswith("heading"):
            level = int(re.sub(r"\D", "", style) or 1)
            sections.append({"level": level, "title": text})
            paragraphs.append(f"\n{'#' * level} {text}")
        else:
            paragraphs.append(text)
    tables = [
        Table(
            header=[cell.text.strip() for cell in table.rows[0].cells] if table.rows else [],
            rows=[[cell.text.strip() for cell in row.cells] for row in table.rows[1:]],
        )
        for table in document.tables
    ]
    return ParsedDocument(
        text="\n".join(paragraphs), tables=tables, sections=sections, format="docx"
    )


# --------------------------------------------------------------------------
# CSV / JSON / TXT
# --------------------------------------------------------------------------

def parse_csv(text: str) -> ParsedDocument:
    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
    except csv.Error:
        dialect = csv.excel
    rows = [row for row in csv.reader(io.StringIO(text), dialect) if any(row)]
    if not rows:
        return ParsedDocument(text="", format="csv")
    header, *body = rows
    return ParsedDocument(text="", tables=[Table(header=header, rows=body)], format="csv")


def parse_json(text: str) -> ParsedDocument:
    payload = json.loads(text)
    tables: list[Table] = []
    if isinstance(payload, list) and payload and isinstance(payload[0], dict):
        header = sorted({key for row in payload for key in row})
        tables.append(
            Table(
                header=header,
                rows=[[str(row.get(k, "")) for k in header] for row in payload],
            )
        )
        body = ""
    else:
        body = json.dumps(payload, ensure_ascii=False, indent=1)
    return ParsedDocument(text=body, tables=tables, format="json")


def parse_txt(text: str) -> ParsedDocument:
    from ai.datasets.cleaning import normalize_encoding

    return ParsedDocument(text=normalize_encoding(text), format="txt")


# --------------------------------------------------------------------------
# Dispatch
# --------------------------------------------------------------------------

def parse_file(path: str | Path) -> ParsedDocument:
    path = Path(path)
    suffix = path.suffix.lower().lstrip(".")
    if suffix in ("html", "htm"):
        return parse_html(path.read_text(encoding="utf-8", errors="replace"))
    if suffix == "pdf":
        return parse_pdf(path.read_bytes())
    if suffix in ("xlsx", "xlsm", "xls"):
        return parse_xlsx(path)
    if suffix == "docx":
        return parse_docx(path)
    if suffix == "csv":
        return parse_csv(path.read_text(encoding="utf-8", errors="replace"))
    if suffix == "json":
        return parse_json(path.read_text(encoding="utf-8", errors="replace"))
    if suffix in ("txt", "md"):
        return parse_txt(path.read_text(encoding="utf-8", errors="replace"))
    raise UnsupportedFormat(f"{path.name}: no parser for .{suffix} (supported: {SUPPORTED})")


__all__ = [
    "ParsedDocument",
    "SUPPORTED",
    "Table",
    "UnsupportedFormat",
    "parse_csv",
    "parse_docx",
    "parse_file",
    "parse_html",
    "parse_json",
    "parse_pdf",
    "parse_txt",
    "parse_xlsx",
]
